import pandas as pd
import re
from datetime import datetime
import os
import zipfile
from pathlib import Path

class ElectionResults:
    def __init__(self, responses_file, valid_numbers_file="valid_numbers.txt"):
        """
        Initialize election results processor
        
        Args:
            responses_file: Path to Google Sheets CSV export
            valid_numbers_file: Path to text file with valid numbers
        """
        print("=" * 80)
        print("FCT COLLEGE OF NURSING SCIENCES ELECTION 2026")
        print("RESULTS PROCESSING SYSTEM")
        print("=" * 80)
        
        # Load valid numbers
        self.valid_numbers = self.load_valid_numbers(valid_numbers_file)
        
        # Load responses
        self.df = self.load_responses(responses_file)
        
        # Setup position mapping
        self.setup_positions()
        
        # Process and validate
        self.validate_votes()
        
        # Get candidate lists (including all candidates, even with 0 votes)
        self.get_all_candidates()
        
        # Calculate results
        self.calculate_results()
    
    def load_valid_numbers(self, filename):
        """Load valid exam numbers from file"""
        if os.path.exists(filename):
            with open(filename, 'r') as f:
                numbers = [line.strip() for line in f if line.strip() and not line.startswith('#')]
            print(f"✓ Loaded {len(numbers)} valid exam numbers from {filename}")
            return numbers
        else:
            print(f"✗ Warning: {filename} not found")
            return []
    
    def load_responses(self, filename):
        """Load election responses"""
        try:
            if filename.endswith('.xlsx'):
                df = pd.read_excel(filename)
            else:
                df = pd.read_csv(filename)
            print(f"✓ Loaded {len(df)} responses from {filename}")
            return df
        except Exception as e:
            print(f"✗ Error loading {filename}: {e}")
            exit(1)
    
    def setup_positions(self):
        """Define all election positions and their columns"""
        self.positions = {
            'PRESIDENT': {
                'column': 'Who are you voting for as President?',
                'candidates': ['EGBUNU VICTOR', 'OLORUNBANWO OLAYINKA PETER']
            },
            'VICE PRESIDENT': {
                'column': 'Who are you voting for as Vice President?',
                'candidates': ['ASONYE GODGIFT NATHAN', 'HALIMAH ISAH']
            },
            'SECRETARY GENERAL': {
                'column': 'Who are you voting for as Secretary General?',
                'candidates': ['OMODARA OLUWABUKOLA FAVOUR', 'ABUBAKAR USMAN (DR SADEEQ)']
            },
            'ASSISTANT SECRETARY GENERAL': {
                'column': 'Who are you voting for as Assistant Secretary General?',
                'candidates': ['EZE JACINTHA CHIDERA', 'SUCCESS ISONG JONAH']
            },
            'TREASURER': {
                'column': 'Who are you voting for as Treasurer?',
                'candidates': ['AHMED KHADIJAT OVAYIOZA', 'EGBELO GLORIA NITEN']
            },
            'FINANCIAL SECRETARY': {
                'column': 'Do you approve ISAAC BULUS MAGODE as Financial Secretary?',
                'type': 'yesno',
                'candidate_name': 'ISAAC BULUS MAGODE'
            },
            'DIRECTOR OF SPECIAL DUTIES': {
                'column': 'Who are you voting for as Director of Special Duties?',
                'candidates': ['UMARATU SUNDAY YUNUSA', 'LUKA ADUWAK ELLEN ZIGWAI', 'ELISHA VICTORIA']
            },
            'DIRECTOR FOR FOOD, WATER & ENVIRONMENTAL SANITATION': {
                'column': 'Who are you voting for as Director for Food, Water and Environmental Sanitation?',
                'candidates': ['OGUNBELE BARNABAS', 'NASIRU ISHAQ']
            },
            'DIRECTOR OF SOCIALS': {
                'column': 'Who are you voting for as Director of Socials?',
                'candidates': ['UMEHAM PRAISE EZINNE', 'ABASIYANGAOWO COLIN JOSEPH']
            },
            'DIRECTOR OF HEALTH': {
                'column': 'Do you approve SHARON CHINDONGNAAN PETER as Director of Health?',
                'type': 'yesno',
                'candidate_name': 'SHARON CHINDONGNAAN PETER'
            },
            'DIRECTOR OF SPORTS': {
                'column': 'Who are you voting for as Director of Sports?',
                'candidates': ['ISAAC PAMBERIMAM PAMELA', 'ISAH USMAN ILIYASU']
            },
            'AUDITOR GENERAL': {
                'column': 'Do you approve JEREMIAH PEACE OMENEFU as Auditor General?',
                'type': 'yesno',
                'candidate_name': 'JEREMIAH PEACE OMENEFU'
            },
            'WELFARE SECRETARY': {
                'column': 'Who are you voting for as Welfare Secretary?',
                'candidates': ['DEHINDO TEMILOLUWA THERESA', 'GEOFFRY DIVINE DESTINY']
            },
            'PRO I': {
                'column': 'Do you approve FATIMA UMAR JEJE as PRO I?',
                'type': 'yesno',
                'candidate_name': 'FATIMA UMAR JEJE'
            },
            'PRO II': {
                'column': 'Do you approve BULUS ALEXANDER FUNOM as PRO II?',
                'type': 'yesno',
                'candidate_name': 'BULUS ALEXANDER FUNOM'
            }
        }
    
    def get_all_candidates(self):
        """Extract all candidates from data (including those with 0 votes)"""
        for position, info in self.positions.items():
            if 'type' not in info or info['type'] != 'yesno':
                column = info['column']
                if column in self.df.columns:
                    # Get unique candidates from data
                    candidates_from_data = self.df[column].dropna().unique()
                    candidates_from_data = [str(c).strip() for c in candidates_from_data]
                    
                    # Combine with predefined list
                    if 'candidates' in info:
                        all_candidates = list(set(info['candidates'] + candidates_from_data))
                    else:
                        all_candidates = candidates_from_data
                    
                    # Sort alphabetically
                    info['candidates'] = sorted(all_candidates)
    
    def validate_votes(self):
        """Validate all votes"""
        print("\n" + "=" * 80)
        print("VALIDATING VOTES")
        print("=" * 80)
        
        # Find exam number column
        exam_cols = [col for col in self.df.columns if any(keyword in col.lower() for keyword in 
                      ['examination', 'exam', 'student', 'number', 'id'])]
        self.exam_column = exam_cols[0] if exam_cols else self.df.columns[2]
        
        print(f"Exam number column: '{self.exam_column}'")
        
        # Clean exam numbers
        self.df['Exam_Clean'] = self.df[self.exam_column].astype(str).str.strip()
        
        # Validate
        self.df['Is_Valid'] = self.df['Exam_Clean'].isin(self.valid_numbers)
        self.df['Is_Duplicate'] = self.df.duplicated(subset=['Exam_Clean'], keep='first')
        self.df['Valid_Vote'] = self.df['Is_Valid'] & (~self.df['Is_Duplicate'])
        
        # Filter valid votes
        self.valid_df = self.df[self.df['Valid_Vote']].copy()
        
        # Statistics
        total = len(self.df)
        valid = self.df['Is_Valid'].sum()
        duplicates = self.df['Is_Duplicate'].sum()
        valid_votes = self.df['Valid_Vote'].sum()
        
        print(f"\n📊 VOTING STATISTICS:")
        print(f"   Total submissions: {total}")
        print(f"   Valid exam numbers: {valid} ({valid/total*100:.1f}%)")
        print(f"   Duplicate votes: {duplicates}")
        print(f"   Valid votes counted: {valid_votes} ({valid_votes/total*100:.1f}%)")
    
    def calculate_results(self):
        """Calculate results for all positions"""
        print("\n" + "=" * 80)
        print("CALCULATING RESULTS")
        print("=" * 80)
        
        self.results = {}
        
        for position, info in self.positions.items():
            column = info['column']
            
            if column not in self.valid_df.columns:
                continue
            
            votes = self.valid_df[column].dropna()
            total_votes = len(votes)
            
            if 'type' in info and info['type'] == 'yesno':
                # Yes/No positions
                yes_votes = votes[votes.astype(str).str.contains('YES', case=False, na=False)].count()
                no_votes = votes[votes.astype(str).str.contains('NO', case=False, na=False)].count()
                
                self.results[position] = {
                    'type': 'yesno',
                    'total': total_votes,
                    'yes': int(yes_votes),
                    'no': int(no_votes),
                    'approval': f"{(yes_votes/total_votes*100):.1f}%" if total_votes > 0 else "0%",
                    'winner': 'YES - CANDIDATE APPROVED' if yes_votes > no_votes else 
                             'NO - CANDIDATE REJECTED' if no_votes > yes_votes else 'TIE - NEEDS REVIEW',
                    'candidate_name': info.get('candidate_name', '')
                }
            else:
                # Multiple candidate positions
                candidates = info.get('candidates', [])
                vote_counts = {}
                
                # Initialize all candidates with 0 votes
                for candidate in candidates:
                    vote_counts[candidate] = 0
                
                # Count actual votes
                actual_counts = votes.value_counts()
                for candidate, count in actual_counts.items():
                    candidate_str = str(candidate).strip()
                    # Find matching candidate (case-insensitive)
                    for known_candidate in candidates:
                        if candidate_str.lower() == known_candidate.lower():
                            vote_counts[known_candidate] = int(count)
                            break
                    else:
                        # If not in list, add it
                        vote_counts[candidate_str] = int(count)
                
                # Calculate percentages
                percentages = {}
                for candidate, count in vote_counts.items():
                    percentages[candidate] = f"{(count/total_votes*100):.1f}%" if total_votes > 0 else "0%"
                
                # Determine winner(s)
                if vote_counts:
                    max_votes = max(vote_counts.values())
                    winners = [candidate for candidate, count in vote_counts.items() if count == max_votes]
                    
                    if len(winners) == 1:
                        winner = winners[0]
                    else:
                        winner = f'TIE BETWEEN: {", ".join(winners)}'
                else:
                    winner = 'NO VALID VOTES'
                
                self.results[position] = {
                    'type': 'candidate',
                    'total': total_votes,
                    'candidates': vote_counts,
                    'percentages': percentages,
                    'winner': winner
                }
    
    def get_signatories_section(self):
        """Return formatted signatories section"""
        return """

OFFICIAL SIGNATORIES
""" + "=" * 80 + """

ISEC CHAIRMAN:      ISTIFANUS BANEDIN VICTOR    _____________________

ISEC SECRETARY:     BABA SABO                   _____________________

SPEAKER:            ABDULHAMID HAJARA           _____________________


Date: _______________
"""
    
    def generate_text_report(self):
        """Generate comprehensive text report with ALL candidates"""
        report_lines = []
        
        # Header
        report_lines.append("=" * 80)
        report_lines.append("FCT COLLEGE OF NURSING SCIENCES, GWAGWALADA")
        report_lines.append("STUDENTS' UNION GOVERNMENT (SUG) ELECTION - 2026")
        report_lines.append("=" * 80)
        report_lines.append(f"Official Election Report - Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report_lines.append("")
        
        # Voting Statistics
        total_submissions = len(self.df)
        valid_votes = len(self.valid_df)
        
        report_lines.append("VOTING STATISTICS")
        report_lines.append("-" * 40)
        report_lines.append(f"Total Submissions            : {total_submissions}")
        report_lines.append(f"Valid Exam Numbers          : {self.df['Is_Valid'].sum()}")
        report_lines.append(f"Duplicate Submissions       : {self.df['Is_Duplicate'].sum()}")
        report_lines.append(f"Valid Votes Counted         : {valid_votes}")
        report_lines.append(f"Invalid Votes (Excluded)    : {total_submissions - valid_votes}")
        report_lines.append(f"Validation Rate             : {(valid_votes/total_submissions*100):.1f}%")
        report_lines.append("")
        
        # Results by Position
        report_lines.append("ELECTION RESULTS BY POSITION")
        report_lines.append("=" * 80)
        report_lines.append("")
        
        # Define display order
        position_order = [
            'PRESIDENT', 'VICE PRESIDENT', 'SECRETARY GENERAL',
            'ASSISTANT SECRETARY GENERAL', 'TREASURER', 'FINANCIAL SECRETARY',
            'DIRECTOR OF SPECIAL DUTIES', 'DIRECTOR FOR FOOD, WATER & ENVIRONMENTAL SANITATION',
            'DIRECTOR OF SOCIALS', 'DIRECTOR OF HEALTH', 'DIRECTOR OF SPORTS',
            'AUDITOR GENERAL', 'WELFARE SECRETARY', 'PRO I', 'PRO II'
        ]
        
        for position in position_order:
            if position in self.results:
                data = self.results[position]
                report_lines.append(position)
                report_lines.append("-" * len(position))
                report_lines.append(f"Total Valid Votes: {data['total']}")
                report_lines.append("")
                
                if data['type'] == 'yesno':
                    # Yes/No positions - show candidate name
                    candidate_name = data.get('candidate_name', 'Unknown Candidate')
                    report_lines.append(f"Candidate: {candidate_name}")
                    report_lines.append("")
                    report_lines.append(f"  YES (Approve)    : {data['yes']:3d} votes ({(data['yes']/data['total']*100):.1f}%)" if data['total'] > 0 else f"  YES (Approve)    : {data['yes']:3d} votes (0.0%)")
                    report_lines.append(f"  NO (Reject)      : {data['no']:3d} votes ({(data['no']/data['total']*100):.1f}%)" if data['total'] > 0 else f"  NO (Reject)      : {data['no']:3d} votes (0.0%)")
                    report_lines.append("")
                    report_lines.append(f"Winner: {data['winner']}")
                else:
                    # Multiple candidates - show ALL candidates including losers
                    report_lines.append("Candidate                                       Votes      Percentage")
                    report_lines.append("-" * 70)
                    
                    # Sort by votes (descending) then alphabetically
                    sorted_candidates = sorted(data['candidates'].items(), 
                                             key=lambda x: (-x[1], x[0]))
                    
                    for candidate, votes in sorted_candidates:
                        percentage = data['percentages'][candidate]
                        report_lines.append(f"{candidate:45}  {votes:3d}        {percentage}")
                    
                    report_lines.append("")
                    report_lines.append(f"Winner: {data['winner']}")
                
                report_lines.append("")
                report_lines.append("")
        
        # Invalid Votes Detail
        report_lines.append("=" * 80)
        report_lines.append("INVALID VOTES DETAIL")
        report_lines.append("-" * 40)
        
        invalid_df = self.df[~self.df['Valid_Vote']]
        if len(invalid_df) > 0:
            invalid_count = (~self.df['Is_Valid']).sum()
            duplicate_count = self.df['Is_Duplicate'].sum()
            
            if invalid_count > 0:
                report_lines.append(f"Invalid Exam Numbers: {invalid_count}")
                invalid_exams = self.df[~self.df['Is_Valid']]['Exam_Clean'].unique()
                report_lines.append(", ".join(invalid_exams[:10]))
                if len(invalid_exams) > 10:
                    report_lines.append(f"... and {len(invalid_exams)-10} more")
            
            if duplicate_count > 0:
                report_lines.append(f"\nDuplicate Votes: {duplicate_count}")
                duplicates = self.df[self.df['Is_Duplicate']]['Exam_Clean'].unique()
                report_lines.append(", ".join(duplicates[:5]))
                if len(duplicates) > 5:
                    report_lines.append(f"... and {len(duplicates)-5} more")
        else:
            report_lines.append("No invalid votes found.")
        
        # Add signatories
        report_lines.append(self.get_signatories_section())
        
        return "\n".join(report_lines)
    
    def generate_word_report(self, filename="election_report.docx"):
        """Generate professional Microsoft Word report with enhanced formatting"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            print(f"\n📝 Generating Word document: {filename}")
            
            doc = Document()
            
            # ====== CUSTOM STYLES ======
            # Add custom styles for professional look
            
            # Title style
            title_style = doc.styles.add_style('ElectionTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
            # Subtitle style
            subtitle_style = doc.styles.add_style('ElectionSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.name = 'Times New Roman'
            subtitle_style.font.size = Pt(18)
            subtitle_style.font.bold = True
            subtitle_style.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(24)
            
            # Section header style
            section_style = doc.styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 0, 0)
            section_style.paragraph_format.space_before = Pt(24)
            section_style.paragraph_format.space_after = Pt(12)
            
            # Position header style
            position_style = doc.styles.add_style('PositionHeader', WD_STYLE_TYPE.PARAGRAPH)
            position_style.font.name = 'Calibri'
            position_style.font.size = Pt(14)
            position_style.font.bold = True
            position_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
            position_style.paragraph_format.space_before = Pt(18)
            position_style.paragraph_format.space_after = Pt(6)
            
            # Normal text style
            normal_style = doc.styles.add_style('ElectionNormal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.5
            
            # Winner style
            winner_style = doc.styles.add_style('WinnerText', WD_STYLE_TYPE.PARAGRAPH)
            winner_style.font.name = 'Calibri'
            winner_style.font.size = Pt(12)
            winner_style.font.bold = True
            winner_style.font.color.rgb = RGBColor(0, 128, 0)  # Green for winners
            
            # ====== HEADER SECTION ======
            # Add logo/header placeholder
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_run = header_para.add_run("FCT COLLEGE OF NURSING SCIENCES, GWAGWALADA")
            header_run.font.name = 'Calibri'
            header_run.font.size = Pt(9)
            header_run.font.italic = True
            header_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Main title
            title = doc.add_paragraph('FCT COLLEGE OF NURSING SCIENCES, GWAGWALADA', style='ElectionTitle')
            
            # Subtitle
            subtitle = doc.add_paragraph("STUDENTS' UNION GOVERNMENT (SUG) ELECTION - 2026", style='ElectionSubtitle')
            
            # Add decorative line
            line = doc.add_paragraph()
            line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = line.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            run.font.bold = True
            run.font.size = Pt(14)
            
            # Report info
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run("OFFICIAL ELECTION RESULTS REPORT")
            info_run.font.name = 'Calibri'
            info_run.font.size = Pt(12)
            info_run.font.bold = True
            info_run.font.color.rgb = RGBColor(128, 0, 0)  # Maroon
            
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", style='ElectionNormal')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # ====== VOTING STATISTICS ======
            doc.add_paragraph('VOTING STATISTICS', style='SectionHeader')
            
            # Create a nicely formatted table for statistics
            stats_table = doc.add_table(rows=6, cols=2)
            stats_table.style = 'LightGrid-Accent1'
            
            # Set table width
            for row in stats_table.rows:
                row.cells[0].width = Cm(8)
                row.cells[1].width = Cm(4)
            
            stats_data = [
                ('Total Submissions', str(len(self.df))),
                ('Valid Exam Numbers', str(self.df['Is_Valid'].sum())),
                ('Duplicate Submissions', str(self.df['Is_Duplicate'].sum())),
                ('Valid Votes Counted', str(len(self.valid_df))),
                ('Invalid Votes (Excluded)', str(len(self.df) - len(self.valid_df))),
                ('Validation Rate', f"{(len(self.valid_df)/len(self.df)*100):.1f}%")
            ]
            
            for i, (label, value) in enumerate(stats_data):
                row = stats_table.rows[i].cells
                row[0].text = label
                row[1].text = value
                
                # Format label cells
                label_para = row[0].paragraphs[0]
                label_para.style = 'ElectionNormal'
                label_run = label_para.runs[0]
                label_run.font.bold = True
                label_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Format value cells
                value_para = row[1].paragraphs[0]
                value_para.style = 'ElectionNormal'
                value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                if 'Rate' in label:
                    value_run = value_para.runs[0]
                    if float(value_para.text.strip('%')) >= 80:
                        value_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for good rates
                    else:
                        value_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for low rates
            
            doc.add_paragraph()
            
            # ====== ELECTION RESULTS ======
            doc.add_paragraph('ELECTION RESULTS BY POSITION', style='SectionHeader')
            
            position_order = [
                'PRESIDENT', 'VICE PRESIDENT', 'SECRETARY GENERAL',
                'ASSISTANT SECRETARY GENERAL', 'TREASURER', 'FINANCIAL SECRETARY',
                'DIRECTOR OF SPECIAL DUTIES', 'DIRECTOR FOR FOOD, WATER & ENVIRONMENTAL SANITATION',
                'DIRECTOR OF SOCIALS', 'DIRECTOR OF HEALTH', 'DIRECTOR OF SPORTS',
                'AUDITOR GENERAL', 'WELFARE SECRETARY', 'PRO I', 'PRO II'
            ]
            
            for position in position_order:
                if position in self.results:
                    data = self.results[position]
                    
                    # Position header
                    doc.add_paragraph(position, style='PositionHeader')
                    
                    # Add a colored background for position header
                    position_para = doc.paragraphs[-1]
                    position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add total votes
                    total_para = doc.add_paragraph(f"Total Valid Votes: {data['total']}", style='ElectionNormal')
                    total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if data['type'] == 'yesno':
                        # Add candidate name
                        candidate_name = data.get('candidate_name', 'Unknown Candidate')
                        candidate_para = doc.add_paragraph(f"Candidate: {candidate_name}", style='ElectionNormal')
                        candidate_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Create results table
                        table = doc.add_table(rows=3, cols=3)
                        table.style = 'LightShading-Accent1'
                        
                        # Set column widths
                        for row in table.rows:
                            row.cells[0].width = Cm(6)
                            row.cells[1].width = Cm(4)
                            row.cells[2].width = Cm(4)
                        
                        # Headers
                        headers = table.rows[0].cells
                        headers[0].text = 'Option'
                        headers[1].text = 'Votes'
                        headers[2].text = 'Percentage'
                        
                        # Format headers
                        for cell in headers:
                            cell.paragraphs[0].runs[0].bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Yes row
                        yes_row = table.rows[1].cells
                        yes_row[0].text = 'YES (Approve)'
                        yes_row[1].text = str(data['yes'])
                        yes_percentage = (data['yes']/data['total']*100) if data['total'] > 0 else 0
                        yes_row[2].text = f"{yes_percentage:.1f}%"
                        
                        # No row
                        no_row = table.rows[2].cells
                        no_row[0].text = 'NO (Reject)'
                        no_row[1].text = str(data['no'])
                        no_percentage = (data['no']/data['total']*100) if data['total'] > 0 else 0
                        no_row[2].text = f"{no_percentage:.1f}%"
                        
                        # Highlight winning row
                        winner_row = yes_row if data['yes'] > data['no'] else no_row if data['no'] > data['yes'] else None
                        if winner_row:
                            for cell in winner_row:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                                        run.font.bold = True
                        
                        # Winner declaration
                        winner_text = data['winner']
                        winner_para = doc.add_paragraph(f"Result: {winner_text}", style='WinnerText')
                        winner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    else:
                        # Multiple candidates - create table
                        num_candidates = len(data['candidates'])
                        table = doc.add_table(rows=num_candidates + 1, cols=3)
                        table.style = 'LightGrid-Accent1'
                        
                        # Set column widths
                        for row in table.rows:
                            row.cells[0].width = Cm(10)
                            row.cells[1].width = Cm(4)
                            row.cells[2].width = Cm(4)
                        
                        # Headers
                        headers = table.rows[0].cells
                        headers[0].text = 'Candidate'
                        headers[1].text = 'Votes'
                        headers[2].text = 'Percentage'
                        
                        # Format headers
                        for cell in headers:
                            cell.paragraphs[0].runs[0].bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.paragraphs[0].style = 'ElectionNormal'
                        
                        # Sort candidates by votes
                        sorted_candidates = sorted(data['candidates'].items(), 
                                                 key=lambda x: (-x[1], x[0]))
                        
                        max_votes = max(data['candidates'].values()) if data['candidates'] else 0
                        
                        for i, (candidate, votes) in enumerate(sorted_candidates, 1):
                            row = table.rows[i].cells
                            row[0].text = candidate
                            row[1].text = str(votes)
                            row[2].text = data['percentages'][candidate]
                            
                            # Center align vote count and percentage
                            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Highlight winner(s)
                            if votes == max_votes and max_votes > 0:
                                # Green background for winner
                                for cell in row:
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), 'E6FFE6')  # Light green
                                    cell._element.tcPr.append(shading_elm)
                                    
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
                                            run.font.bold = True
                        
                        # Winner declaration
                        winner_para = doc.add_paragraph(f"Winner: {data['winner']}", style='WinnerText')
                        winner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add spacing between positions
                    doc.add_paragraph()
            
            # ====== INVALID VOTES SECTION ======
            doc.add_page_break()
            doc.add_paragraph('INVALID VOTES DETAIL', style='SectionHeader')
            
            invalid_df = self.df[~self.df['Valid_Vote']]
            
            if len(invalid_df) > 0:
                invalid_count = (~self.df['Is_Valid']).sum()
                duplicate_count = self.df['Is_Duplicate'].sum()
                
                if invalid_count > 0:
                    doc.add_paragraph('Invalid Exam Numbers', style='PositionHeader')
                    count_para = doc.add_paragraph(f"Total: {invalid_count}", style='ElectionNormal')
                    
                    invalid_exams = self.df[~self.df['Is_Valid']]['Exam_Clean'].unique()
                    exams_para = doc.add_paragraph(", ".join(invalid_exams), style='ElectionNormal')
                    exams_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                if duplicate_count > 0:
                    doc.add_paragraph('Duplicate Votes', style='PositionHeader')
                    dup_para = doc.add_paragraph(f"Total: {duplicate_count}", style='ElectionNormal')
                    
                    duplicates = self.df[self.df['Is_Duplicate']]['Exam_Clean'].unique()
                    dup_list = doc.add_paragraph(", ".join(duplicates), style='ElectionNormal')
                    dup_list.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                doc.add_paragraph("No invalid votes found.", style='ElectionNormal')
            
            # ====== OFFICIAL SIGNATORIES SECTION ======
            doc.add_page_break()
            
            # Signatories header
            sig_header = doc.add_paragraph('OFFICIAL SIGNATORIES', style='SectionHeader')
            sig_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add decorative border
            border_para = doc.add_paragraph()
            border_run = border_para.add_run("═" * 60)
            border_run.font.color.rgb = RGBColor(0, 51, 102)
            border_run.font.bold = True
            border_run.font.size = Pt(12)
            border_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add note
            note = doc.add_paragraph("This document serves as the official record of the 2026 SUG Election results.", style='ElectionNormal')
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # ====== PROFESSIONAL SIGNATORIES LAYOUT ======
            # Add spacing
            for _ in range(3):
                doc.add_paragraph()
            
            # ISEC CHAIRMAN
            doc.add_paragraph('ISEC CHAIRMAN', style='PositionHeader')
            
            chairman_name = doc.add_paragraph('ISTIFANUS BANEDIN VICTOR', style='ElectionNormal')
            chairman_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Signature line
            sig_line1 = doc.add_paragraph()
            sig_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_run1 = sig_line1.add_run("_" * 40)
            sig_run1.font.size = Pt(14)
            
            sig_text1 = doc.add_paragraph('Signature', style='ElectionNormal')
            sig_text1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date line
            date_line1 = doc.add_paragraph()
            date_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run1 = date_line1.add_run("Date: _______________")
            date_run1.font.size = Pt(11)
            
            # Add spacing
            for _ in range(2):
                doc.add_paragraph()
            
            # ISEC SECRETARY
            doc.add_paragraph('ISEC SECRETARY', style='PositionHeader')
            
            secretary_name = doc.add_paragraph('BABA SABO', style='ElectionNormal')
            secretary_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Signature line
            sig_line2 = doc.add_paragraph()
            sig_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_run2 = sig_line2.add_run("_" * 40)
            sig_run2.font.size = Pt(14)
            
            sig_text2 = doc.add_paragraph('Signature', style='ElectionNormal')
            sig_text2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date line
            date_line2 = doc.add_paragraph()
            date_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run2 = date_line2.add_run("Date: _______________")
            date_run2.font.size = Pt(11)
            
            # Add spacing
            for _ in range(2):
                doc.add_paragraph()
            
            # SPEAKER
            doc.add_paragraph('SPEAKER', style='PositionHeader')
            
            speaker_name = doc.add_paragraph('ABDULHAMID HAJARA', style='ElectionNormal')
            speaker_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Signature line
            sig_line3 = doc.add_paragraph()
            sig_line3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_run3 = sig_line3.add_run("_" * 40)
            sig_run3.font.size = Pt(14)
            
            sig_text3 = doc.add_paragraph('Signature', style='ElectionNormal')
            sig_text3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date line
            date_line3 = doc.add_paragraph()
            date_line3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run3 = date_line3.add_run("Date: _______________")
            date_run3.font.size = Pt(11)
            
            # Final decorative border
            doc.add_paragraph()
            final_border = doc.add_paragraph()
            final_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
            final_run = final_border.add_run("═" * 60)
            final_run.font.color.rgb = RGBColor(0, 51, 102)
            final_run.font.bold = True
            final_run.font.size = Pt(12)
            
            # Footer with page numbers
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(f"FCT College of Nursing Sciences - SUG Election 2026 - Page ")
            
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            footer_run._element.append(fldChar1)
            footer_run._element.append(instrText)
            footer_run._element.append(fldChar2)
            
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Save document
            doc.save(filename)
            print(f"✓ Word report saved: {filename}")
            return filename
            
        except ImportError:
            print("✗ python-docx not installed. Installing...")
            import subprocess
            subprocess.check_call(['pip', 'install', 'python-docx'])
            return self.generate_word_report(filename)
    
    def export_to_excel(self, filename="election_results.xlsx"):
        """Export detailed results to Excel with invalid votes sheet"""
        print(f"\n📊 Generating Excel report: {filename}")
        
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            # Summary sheet
            summary_data = {
                'Metric': [
                    'Total Submissions',
                    'Valid Exam Numbers',
                    'Duplicate Submissions',
                    'Valid Votes Counted',
                    'Invalid Votes (Excluded)',
                    'Validation Rate'
                ],
                'Value': [
                    len(self.df),
                    self.df['Is_Valid'].sum(),
                    self.df['Is_Duplicate'].sum(),
                    len(self.valid_df),
                    len(self.df) - len(self.valid_df),
                    f"{(len(self.valid_df)/len(self.df)*100):.1f}%"
                ]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Results sheet - WITH ALL CANDIDATES
            results_data = []
            for position, data in self.results.items():
                if data['type'] == 'yesno':
                    candidate_name = data.get('candidate_name', 'Unknown')
                    results_data.append({
                        'Position': position,
                        'Candidate': f"{candidate_name} - YES (Approve)",
                        'Votes': data['yes'],
                        'Percentage': f"{(data['yes']/data['total']*100):.1f}%" if data['total'] > 0 else "0%",
                        'Status': 'APPROVED' if data['yes'] > data['no'] else ''
                    })
                    results_data.append({
                        'Position': position,
                        'Candidate': f"{candidate_name} - NO (Reject)",
                        'Votes': data['no'],
                        'Percentage': f"{(data['no']/data['total']*100):.1f}%" if data['total'] > 0 else "0%",
                        'Status': 'REJECTED' if data['no'] > data['yes'] else ''
                    })
                else:
                    # Sort candidates by votes for better readability
                    sorted_candidates = sorted(data['candidates'].items(), 
                                             key=lambda x: (-x[1], x[0]))
                    for candidate, votes in sorted_candidates:
                        results_data.append({
                            'Position': position,
                            'Candidate': candidate,
                            'Votes': votes,
                            'Percentage': data['percentages'][candidate],
                            'Status': 'ELECTED' if candidate in str(data['winner']) else ''
                        })
            
            results_df = pd.DataFrame(results_data)
            results_df.to_excel(writer, sheet_name='Results', index=False)
            
            # Invalid votes sheet
            invalid_df = self.df[~self.df['Valid_Vote']].copy()
            invalid_df['Reason'] = invalid_df.apply(
                lambda row: 'Invalid Exam Number' if not row['Is_Valid'] 
                           else 'Duplicate Vote' if row['Is_Duplicate'] 
                           else 'Unknown', axis=1
            )
            # Select relevant columns
            invalid_cols = ['Exam_Clean', 'Reason'] + [col for col in invalid_df.columns 
                           if col not in ['Exam_Clean', 'Reason', 'Is_Valid', 'Is_Duplicate', 'Valid_Vote']]
            invalid_df[invalid_cols].to_excel(writer, sheet_name='Invalid_Votes', index=False)
            
            # Valid votes sheet
            self.valid_df.to_excel(writer, sheet_name='Valid_Votes', index=False)
            
            # All data sheet
            self.df.to_excel(writer, sheet_name='All_Data', index=False)
        
        print(f"✓ Excel report saved: {filename}")
        print(f"  - Summary sheet with voting statistics")
        print(f"  - Results sheet with ALL candidates (winners and losers)")
        print(f"  - Invalid_Votes sheet with {len(self.df[~self.df['Valid_Vote']])} invalid submissions")
        print(f"  - Valid_Votes sheet with {len(self.valid_df)} valid votes")
        print(f"  - All_Data sheet with complete dataset")
        return filename
    
    def create_zip_package(self, output_dir="election_results"):
        """Package all reports into a ZIP file"""
        print("\n" + "=" * 80)
        print("CREATING ZIP PACKAGE")
        print("=" * 80)
        
        # Create output directory
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Generate all reports in the output directory
        text_file = os.path.join(output_dir, "election_report.txt")
        word_file = os.path.join(output_dir, "election_report.docx")
        excel_file = os.path.join(output_dir, "election_results.xlsx")
        
        # Generate text report
        text_report = self.generate_text_report()
        with open(text_file, "w", encoding='utf-8') as f:
            f.write(text_report)
        print(f"✓ Text report saved: {text_file}")
        
        # Generate Word report
        self.generate_word_report(word_file)
        
        # Generate Excel report
        self.export_to_excel(excel_file)
        
        # Create ZIP file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f"election_results_{timestamp}.zip"
        
        print(f"\n📦 Creating ZIP archive: {zip_filename}")
        
        with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add all files from output directory
            for root, dirs, files in os.walk(output_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, output_dir)
                    zipf.write(file_path, arcname)
                    print(f"  + Added: {arcname}")
        
        print(f"\n✓ ZIP package created successfully: {zip_filename}")
        print(f"  Size: {os.path.getsize(zip_filename) / 1024:.2f} KB")
        
        # Clean up temporary directory
        import shutil
        shutil.rmtree(output_dir)
        print(f"✓ Cleaned up temporary files")
        
        return zip_filename

def main():
    print("=" * 80)
    print("FCT COLLEGE OF NURSING SCIENCES - ELECTION 2026")
    print("COMPREHENSIVE RESULTS PROCESSOR")
    print("=" * 80)
    
    # Check files
    responses_file = "election_responses.csv"
    valid_file = "valid_numbers.txt"
    
    if not os.path.exists(responses_file):
        print(f"✗ ERROR: {responses_file} not found!")
        print(f"Current directory: {os.getcwd()}")
        print("\nAvailable CSV/Excel files:")
        for f in os.listdir('.'):
            if f.endswith(('.csv', '.xlsx')):
                print(f"  - {f}")
        return
    
    if not os.path.exists(valid_file):
        print(f"⚠ WARNING: {valid_file} not found")
        print("Creating empty valid numbers file...")
        with open(valid_file, 'w') as f:
            f.write("# Add valid exam numbers here, one per line\n")
            f.write("# Example:\n")
            f.write("# 6602\n")
            f.write("# 6671\n")
            f.write("# FCTCONS/ND23/001\n")
    
    # Process election
    print("\n" + "=" * 80)
    print("PROCESSING ELECTION DATA")
    print("=" * 80)
    
    election = ElectionResults(responses_file, valid_file)
    
    # Generate reports and create ZIP package
    print("\n" + "=" * 80)
    print("GENERATING REPORTS")
    print("=" * 80)
    
    zip_file = election.create_zip_package()
    
    # Display summary
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE!")
    print("=" * 80)
    print(f"📦 All reports packaged in: {zip_file}")
    print(f"\nPackage contents:")
    print(f"  1. election_report.txt - Full text report with signatories")
    print(f"  2. election_report.docx - Word document with signatories")
    print(f"  3. election_results.xlsx - Excel workbook with multiple sheets")
    print(f"\nElection Statistics:")
    print(f"  • Valid Votes Processed: {len(election.valid_df)}")
    print(f"  • Invalid Votes Excluded: {len(election.df) - len(election.valid_df)}")
    print(f"  • Validation Rate: {(len(election.valid_df)/len(election.df)*100):.1f}%")
    print("=" * 80)
    
    # Print summary to console
    print("\n📋 ELECTION WINNERS SUMMARY:")
    print("-" * 80)
    for position, data in election.results.items():
        if data['type'] == 'yesno':
            status = "✓ APPROVED" if data['yes'] > data['no'] else "✗ REJECTED"
            print(f"  {position}:")
            print(f"    {data.get('candidate_name', 'Unknown')} - {status} ({data['approval']})")
        else:
            winner = data['winner'].split(':')[-1].strip() if 'TIE' in data['winner'] else data['winner']
            print(f"  {position}: {winner}")
    print("=" * 80)

if __name__ == "__main__":
    # Check for required packages
    required = ['pandas']
    missing = []
    
    for package in required:
        try:
            __import__(package)
        except ImportError:
            missing.append(package)
    
    if missing:
        print(f"Installing required packages: {', '.join(missing)}")
        import subprocess
        subprocess.check_call(['pip', 'install'] + missing + ['openpyxl'])
        print("Packages installed. Please run the script again.")
    else:
        main()